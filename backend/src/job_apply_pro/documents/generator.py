from __future__ import annotations

from io import BytesIO

from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Flowable, Paragraph, SimpleDocTemplate, Spacer

from job_apply_pro.domain.knowledge import (
    DocumentOutputFormat,
    DocumentTemplate,
    TailoredDocumentPreview,
)


def render_tailored_document(preview: TailoredDocumentPreview) -> tuple[str, bytes]:
    stem = "-".join(part for part in (preview.title, preview.employer, preview.kind.value) if part)
    safe_stem = "".join(
        character if character.isalnum() or character in {"-", "_"} else "-" for character in stem
    ).strip("-")[:100]
    if preview.output_format is DocumentOutputFormat.DOCX:
        return f"{safe_stem}.docx", _render_docx(preview)
    return f"{safe_stem}.pdf", _render_pdf(preview)


def _render_docx(preview: TailoredDocumentPreview) -> bytes:
    document = DocxDocument()
    document.core_properties.title = f"{preview.title} - {preview.employer}"
    if preview.template is DocumentTemplate.COMPACT:
        for document_section in document.sections:
            document_section.top_margin = Inches(0.45)
            document_section.bottom_margin = Inches(0.45)
            document_section.left_margin = Inches(0.55)
            document_section.right_margin = Inches(0.55)
        document.styles["Normal"].font.size = Pt(9)
    else:
        document.styles["Normal"].font.size = Pt(10.5)
    for section_index, section in enumerate(preview.sections):
        document.add_heading(section.heading, level=1 if section_index == 0 else 2)
        for paragraph in section.paragraphs:
            document.add_paragraph(paragraph)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _render_pdf(preview: TailoredDocumentPreview) -> bytes:
    buffer = BytesIO()
    styles = getSampleStyleSheet()
    if preview.template is DocumentTemplate.COMPACT:
        styles["BodyText"].fontSize = 8.5
        styles["BodyText"].leading = 10
        styles["Heading2"].spaceBefore = 4
        styles["Heading2"].spaceAfter = 2
    story: list[Flowable] = []
    for section_index, section in enumerate(preview.sections):
        style = styles["Title"] if section_index == 0 else styles["Heading2"]
        story.append(Paragraph(section.heading, style))
        story.append(Spacer(1, 8))
        for text in section.paragraphs:
            story.append(
                Paragraph(text.replace("&", "&amp;").replace("<", "&lt;"), styles["BodyText"])
            )
            story.append(Spacer(1, 3 if preview.template is DocumentTemplate.COMPACT else 6))
    margin = 32 if preview.template is DocumentTemplate.COMPACT else 54
    SimpleDocTemplate(
        buffer,
        pagesize=LETTER,
        title=preview.title,
        leftMargin=margin,
        rightMargin=margin,
        topMargin=margin,
        bottomMargin=margin,
    ).build(story)
    return buffer.getvalue()
