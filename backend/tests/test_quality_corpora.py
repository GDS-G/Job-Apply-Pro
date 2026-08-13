from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path
from typing import cast

import pytest
from docx import Document as DocxDocument
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from reportlab.pdfgen.canvas import Canvas
from sqlalchemy.orm import Session

from job_apply_pro.ai.prompts import AGENT_TASKS
from job_apply_pro.documents.extractors import extract_document
from job_apply_pro.domain.ai import AIRoutingPolicy
from job_apply_pro.quality.corpora import (
    AIEvaluationCorpus,
    CorpusValidationError,
    DocumentFixtureKind,
    DocumentIngestionCorpusCase,
    parse_ai_evaluation_corpus,
    parse_document_ingestion_corpus,
    run_document_ingestion_corpus,
)
from job_apply_pro.services.ai import AgentService, AIEvaluationHarness
from test_ai_gateway import FakeProvider, _models, _policy, _service

FIXTURES = Path(__file__).parent / "fixtures"


def _load_ai_corpus() -> AIEvaluationCorpus:
    return parse_ai_evaluation_corpus((FIXTURES / "ai_evaluation_corpus_v1.json").read_bytes())


def _fixture_bytes(case: DocumentIngestionCorpusCase) -> bytes:
    values = case.fixture_values
    if case.fixture_kind in {DocumentFixtureKind.PLAIN_TEXT, DocumentFixtureKind.MARKDOWN}:
        return "\n".join(str(value) for value in cast(list[object], values["lines"])).encode()
    if case.fixture_kind == DocumentFixtureKind.RTF:
        paragraphs = r"\par ".join(str(value) for value in cast(list[object], values["paragraphs"]))
        return (r"{\rtf1\ansi " + paragraphs + "}").encode()
    if case.fixture_kind in {
        DocumentFixtureKind.DOCX_SECTION_TABLE,
        DocumentFixtureKind.DOCX_NESTED_FLOATING,
    }:
        output = BytesIO()
        document = DocxDocument()
        if case.fixture_kind == DocumentFixtureKind.DOCX_SECTION_TABLE:
            document.add_paragraph(str(values["heading"]))
            table = document.add_table(rows=1, cols=2)
            cells = cast(list[object], values["cells"])
            table.cell(0, 0).text = str(cells[0])
            table.cell(0, 1).text = str(cells[1])
            document.add_paragraph(str(values["ending"]))
        else:
            outer = document.add_table(rows=1, cols=1)
            outer.cell(0, 0).paragraphs[0].add_run(str(values["outer"]))
            nested = outer.cell(0, 0).add_table(rows=1, cols=2)
            nested_values = cast(list[object], values["nested"])
            nested.cell(0, 0).text = str(nested_values[0])
            nested.cell(0, 1).text = str(nested_values[1])
            floating = document.add_paragraph()
            floating._p.append(
                parse_xml(
                    f"""<w:pict {nsdecls("w")}><w:txbxContent><w:p><w:r>
                    <w:t>{values["floating"]}</w:t></w:r></w:p>
                    </w:txbxContent></w:pict>"""
                )
            )
        document.save(output)
        return output.getvalue()
    output = BytesIO()
    canvas = Canvas(output)
    if case.fixture_kind == DocumentFixtureKind.PDF_SINGLE_COLUMN:
        for index, line in enumerate(cast(list[object], values["lines"])):
            canvas.drawString(72, 740 - index * 24, str(line))
    elif case.fixture_kind == DocumentFixtureKind.PDF_TWO_COLUMN:
        canvas.drawString(72, 740, str(values["heading"]))
        for index, line in enumerate(cast(list[object], values["left"])):
            canvas.drawString(72, 700 - index * 20, str(line))
        for index, line in enumerate(cast(list[object], values["right"])):
            canvas.drawString(330, 700 - index * 20, str(line))
    elif case.fixture_kind == DocumentFixtureKind.PDF_PARTIAL_BLANK:
        canvas.drawString(72, 740, str(values["text"]))
        canvas.showPage()
        canvas.showPage()
    else:  # pragma: no cover - all enum variants are handled above
        raise AssertionError(f"Unhandled fixture kind: {case.fixture_kind}")
    canvas.save()
    return output.getvalue()


def test_versioned_ai_corpus_runs_every_agent_role_twice(session: Session) -> None:
    corpus = _load_ai_corpus()
    outputs: list[str | Exception] = [
        json.dumps(case.fixture_output) for case in corpus.cases for _ in range(2)
    ]
    provider = FakeProvider("local", outputs)
    models = _models("local")
    policies: list[AIRoutingPolicy] = []
    for task in set(AGENT_TASKS.values()):
        policy = _policy(models[0].id).model_copy(update={"task_type": task})
        policies.append(policy)
    report = AIEvaluationHarness(AgentService(_service(session, [provider], models, policies))).run(
        corpus.evaluation_cases()
    )

    assert {case.role for case in corpus.cases} == set(AGENT_TASKS)
    assert report.total == report.passed == 6
    assert report.failed == 0
    assert all(len(result.invocation_ids) == 2 for result in report.cases)
    assert all(result.output_fingerprint for result in report.cases)
    assert len(provider.requests) == len(corpus.cases) * 2


def test_versioned_document_corpus_matches_extractors() -> None:
    corpus = parse_document_ingestion_corpus(
        (FIXTURES / "document_ingestion_corpus_v1.json").read_bytes()
    )
    report = run_document_ingestion_corpus(corpus, _fixture_bytes, extract_document)

    assert report.corpus_version == "1.0.0"
    assert report.total == report.passed == 8
    assert report.failed == 0


@pytest.mark.parametrize(
    "mutation",
    [
        lambda payload: payload["cases"].append(payload["cases"][0]),
        lambda payload: payload["cases"][0]["input_data"].update(
            {"password": "credential-shaped-value"}
        ),
        lambda payload: payload["cases"][0]["fixture_output"].update(
            {"reason": "person@example.test"}
        ),
    ],
)
def test_ai_corpus_rejects_duplicate_or_sensitive_fixtures(mutation: object) -> None:
    payload = json.loads((FIXTURES / "ai_evaluation_corpus_v1.json").read_text())
    mutation(payload)  # type: ignore[operator]
    with pytest.raises(CorpusValidationError, match="failed validation"):
        parse_ai_evaluation_corpus(json.dumps(payload).encode())


def test_corpus_loader_rejects_oversized_or_invalid_json() -> None:
    with pytest.raises(CorpusValidationError, match="size limit"):
        parse_ai_evaluation_corpus(b"x" * 1_000_001)
    with pytest.raises(CorpusValidationError, match="failed validation"):
        parse_ai_evaluation_corpus(b"not-json")


def test_document_corpus_rejects_mismatched_expected_blocks() -> None:
    payload = json.loads((FIXTURES / "document_ingestion_corpus_v1.json").read_text())
    payload["cases"][0]["expected_block_kinds"] = ["LINE"]
    with pytest.raises(CorpusValidationError, match="failed validation"):
        parse_document_ingestion_corpus(json.dumps(payload).encode())
