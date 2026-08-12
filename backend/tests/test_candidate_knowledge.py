from __future__ import annotations

import hashlib
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from reportlab.pdfgen.canvas import Canvas
from sqlalchemy import select
from sqlalchemy.orm import Session

from job_apply_pro.api.routes.knowledge import get_knowledge_service
from job_apply_pro.documents.extractors import extract_document
from job_apply_pro.domain.ai import AIRerankRequest, AIRerankResult, DataClassification
from job_apply_pro.domain.applications import (
    ApplicationCreate,
    ApplicationDocumentRole,
    SubmittedDocumentCapture,
)
from job_apply_pro.domain.candidate import CandidateProfileCreate, ContactDetails
from job_apply_pro.domain.jobs import JobCreate
from job_apply_pro.domain.knowledge import (
    AnswerLibraryCreate,
    AnswerLibraryUpdate,
    ClaimPermittedUse,
    ClaimReview,
    ClaimVerificationStatus,
    DocumentKind,
    DocumentOutputFormat,
    DocumentSelectionApproval,
    DocumentSelectionRequest,
    DocumentTemplate,
    RetrievalQuery,
    TailoredDocumentApproval,
    TailoredDocumentRequest,
    TailoringRankingMode,
)
from job_apply_pro.main import app
from job_apply_pro.security.encryption import SensitiveDataCipher
from job_apply_pro.security.keys import StaticKeyProvider
from job_apply_pro.services.core import CoreService
from job_apply_pro.services.knowledge import (
    CandidateKnowledgeConflictError,
    CandidateKnowledgeService,
)
from job_apply_pro.storage.knowledge_repository import CandidateKnowledgeRepository
from job_apply_pro.storage.models import (
    AnswerLibraryRevisionRow,
    AnswerLibraryRow,
    DocumentGenerationAuditRow,
    DocumentVersionRow,
    JobRequirementRow,
    RetrievalChunkRow,
)
from job_apply_pro.storage.repositories import (
    ApplicationRepository,
    CandidateRepository,
    CheckpointRepository,
    JobRepository,
)


def _profile(session: Session) -> str:
    candidates = CandidateRepository(session)
    profile = CoreService(
        candidates,
        JobRepository(session),
        ApplicationRepository(session),
        CheckpointRepository(session),
        SensitiveDataCipher(StaticKeyProvider(b"k" * 32)),
    ).create_candidate(
        CandidateProfileCreate(
            display_name="Knowledge profile",
            contact=ContactDetails(
                full_name="Knowledge User",
                email="knowledge@example.com",
            ),
        )
    )
    return profile.id


def _service(session: Session, tmp_path: Path) -> CandidateKnowledgeService:
    return CandidateKnowledgeService(
        CandidateKnowledgeRepository(session),
        CandidateRepository(session),
        JobRepository(session),
        ApplicationRepository(session),
        SensitiveDataCipher(StaticKeyProvider(b"k" * 32)),
        document_data_dir=tmp_path / "documents",
        document_max_bytes=1_000_000,
    )


RESUME = b"""Knowledge User
knowledge.user@example.com | (312) 555-0142
Senior Network Engineer | Jan 2020 - Mar 2023 | BGP automation
Network Engineer | Jan 2022 - Dec 2024 | BGP Python automation
Skills: Python, BGP, automation, Docker, SQL
Certification: CCNP
"""


def test_import_review_lock_experience_and_private_retrieval(
    session: Session, tmp_path: Path
) -> None:
    profile_id = _profile(session)
    service = _service(session, tmp_path)
    imported = service.import_document(
        profile_id,
        file_name="network-resume.txt",
        data=RESUME,
        kind=DocumentKind.RESUME,
        display_name="Network resume",
        variant_label="Network engineering",
        job_family_tags=["network", "infrastructure", "network"],
        is_primary=True,
    )

    assert imported.document.variant_label == "Network engineering"
    assert imported.document.job_family_tags == ["infrastructure", "network"]
    assert imported.extraction.character_count > 100
    assert all(
        claim.verification_status is ClaimVerificationStatus.PROPOSED
        for claim in imported.proposed_claims
    )
    stored_version = session.scalar(
        select(DocumentVersionRow).where(DocumentVersionRow.id == imported.version.id)
    )
    assert stored_version is not None
    assert "knowledge.user@example.com" not in stored_version.encrypted_extraction
    assert "knowledge.user@example.com" not in Path(stored_version.storage_path).read_text()
    assert service.get_extraction(imported.version.id) == imported.extraction

    bgp_skill = next(
        claim for claim in imported.proposed_claims if claim.canonical_key == "skill.bgp"
    )
    verified_skill = service.review_claim(
        bgp_skill.id,
        ClaimReview(
            approved=True,
            permitted_use=ClaimPermittedUse.APPLICATIONS,
        ),
    )
    assert verified_skill.locked
    experience_claims = [
        claim
        for claim in imported.proposed_claims
        if claim.canonical_key.startswith("experience.bgp.")
    ]
    for claim in experience_claims:
        service.review_claim(
            claim.id,
            ClaimReview(
                approved=True,
                permitted_use=ClaimPermittedUse.APPLICATIONS,
            ),
        )
    summary = next(item for item in service.calculate_experience(profile_id) if item.skill == "bgp")
    assert summary.months == 60
    assert summary.years == 5

    answer = service.add_answer(
        profile_id,
        AnswerLibraryCreate(
            question="How many years of BGP experience do you have?",
            canonical_field="experience.bgp_years",
            answer="Five years of verified BGP experience.",
            evidence_claim_ids=[verified_skill.id, *[claim.id for claim in experience_claims]],
            provenance={"reviewed_by": "user"},
            confirmation_phrase="SAVE REVIEWED ANSWER",
        ),
    )
    assert answer.revision == 1
    assert answer.locked and answer.approved
    results = service.retrieve(
        profile_id,
        RetrievalQuery(query="BGP experience", permitted_use=ClaimPermittedUse.APPLICATIONS),
    )
    assert {result.source_type for result in results} == {"ANSWER", "CLAIM"}
    assert any("Five years" in result.content for result in results)

    answer_row = session.get(AnswerLibraryRow, answer.id)
    assert answer_row is not None
    assert "Five years" not in answer_row.encrypted_answer
    chunks = session.scalars(select(RetrievalChunkRow)).all()
    assert chunks
    assert all("bgp" not in " ".join(chunk.token_hashes_json).casefold() for chunk in chunks)
    cipher = SensitiveDataCipher(StaticKeyProvider(b"k" * 32))
    decrypted_chunks = [
        cipher.decrypt_bytes(
            chunk.encrypted_content,
            context=f"retrieval:{chunk.id}",
        ).decode()
        for chunk in chunks
    ]
    assert all(chunk.encrypted_content.startswith("jap:v1:") for chunk in chunks)
    assert all(
        chunk.encrypted_content != decrypted
        for chunk, decrypted in zip(chunks, decrypted_chunks, strict=True)
    )
    assert any("BGP" in decrypted for decrypted in decrypted_chunks)
    assert any("Five years" in decrypted for decrypted in decrypted_chunks)

    corrected = service.update_answer(
        answer.id,
        AnswerLibraryUpdate(
            question="How many years of BGP experience do you have?",
            canonical_field="experience.bgp_years",
            answer="I have five years of verified BGP experience.",
            evidence_claim_ids=answer.evidence_claim_ids,
            provenance={"reviewed_by": "user", "reason": "clarity"},
            expected_revision=answer.revision,
            confirmation_phrase="SAVE REVIEWED ANSWER",
        ),
    )
    assert corrected.revision == 2
    assert corrected.answer.startswith("I have")
    revisions = service.list_answer_revisions(answer.id)
    assert [revision.revision for revision in revisions] == [2, 1]
    assert revisions[1].answer == "Five years of verified BGP experience."
    revision_rows = session.scalars(select(AnswerLibraryRevisionRow)).all()
    assert len(revision_rows) == 2
    assert all("Five years" not in row.encrypted_answer for row in revision_rows)
    with pytest.raises(CandidateKnowledgeConflictError, match="refresh"):
        service.update_answer(
            answer.id,
            AnswerLibraryUpdate(
                question=answer.question,
                canonical_field=answer.canonical_field,
                answer="A stale correction must not win.",
                evidence_claim_ids=answer.evidence_claim_ids,
                expected_revision=1,
                confirmation_phrase="SAVE REVIEWED ANSWER",
            ),
        )
    refreshed_results = service.retrieve(
        profile_id,
        RetrievalQuery(query="BGP experience", permitted_use=ClaimPermittedUse.APPLICATIONS),
    )
    answer_result = next(result for result in refreshed_results if result.source_type == "ANSWER")
    assert "I have five years" in answer_result.content
    assert "\nFive years" not in answer_result.content

    disabled = service.update_answer(
        answer.id,
        AnswerLibraryUpdate(
            question=corrected.question,
            canonical_field=corrected.canonical_field,
            answer=corrected.answer,
            evidence_claim_ids=corrected.evidence_claim_ids,
            approved=False,
            locked=False,
            expected_revision=corrected.revision,
            confirmation_phrase="SAVE REVIEWED ANSWER",
        ),
    )
    assert disabled.revision == 3
    assert not disabled.approved and not disabled.locked
    disabled_results = service.retrieve(
        profile_id,
        RetrievalQuery(query="BGP experience", permitted_use=ClaimPermittedUse.APPLICATIONS),
    )
    assert all(result.source_type != "ANSWER" for result in disabled_results)

    duplicate = service.import_document(
        profile_id,
        file_name="network-resume-v2.txt",
        data=RESUME.replace(b"Docker", b"Kubernetes"),
        kind=DocumentKind.RESUME,
        display_name="Network resume v2",
        variant_label="Network engineering alternate",
        job_family_tags=["network"],
        is_primary=False,
    )
    duplicate_bgp = next(
        claim for claim in duplicate.proposed_claims if claim.canonical_key == "skill.bgp"
    )
    with pytest.raises(CandidateKnowledgeConflictError, match="locked verified fact"):
        service.review_claim(duplicate_bgp.id, ClaimReview(approved=True))
    assert service.list_claims(profile_id)[-1].verification_status in {
        ClaimVerificationStatus.PROPOSED,
        ClaimVerificationStatus.VERIFIED,
    }


def test_explainable_document_selection_requires_review_and_persists_audit(
    session: Session, tmp_path: Path
) -> None:
    profile_id = _profile(session)
    service = _service(session, tmp_path)
    platform = service.import_document(
        profile_id,
        file_name="platform.txt",
        data=b"Platform Engineer\nKubernetes Python operations and distributed systems",
        kind=DocumentKind.RESUME,
        display_name="Platform resume",
        variant_label="Platform engineering",
        job_family_tags=["platform", "cloud"],
        is_primary=False,
    )
    service.import_document(
        profile_id,
        file_name="general.txt",
        data=b"Generalist\nCustomer support and office administration",
        kind=DocumentKind.RESUME,
        display_name="General resume",
        variant_label="General",
        job_family_tags=["general"],
        is_primary=True,
    )
    job = JobRepository(session).add(
        JobCreate(
            source="fixture",
            external_id="selection-1",
            employer="Evidence Systems",
            title="Senior Platform Engineer",
            description_hash=hashlib.sha256(b"selection").hexdigest(),
        )
    )
    session.add_all(
        [
            JobRequirementRow(
                id="selection-kubernetes",
                job_id=job.id,
                category="skill",
                text="Kubernetes operations",
                required=True,
                evidence_json={"source": "fixture"},
            ),
            JobRequirementRow(
                id="selection-python",
                job_id=job.id,
                category="skill",
                text="Python",
                required=False,
                evidence_json={"source": "fixture"},
            ),
        ]
    )
    session.commit()
    application = ApplicationRepository(session).add(
        ApplicationCreate(
            workflow_id="selection-workflow",
            profile_id=profile_id,
            job_id=job.id,
        )
    )
    request = DocumentSelectionRequest(
        application_id=application.id,
        preferred_tags=["cloud", "platform", "cloud"],
    )

    preview = service.preview_document_selection(request)
    assert preview.recommended_document_version_id == platform.version.id
    assert preview.recommendations[0].document_id == platform.document.id
    assert preview.recommendations[0].matched_job_family_tags == ["cloud", "platform"]
    assert preview.recommendations[0].matched_requirement_ids == [
        "selection-kubernetes",
        "selection-python",
    ]
    assert any("2 of 2" in reason for reason in preview.recommendations[0].reasons)

    with pytest.raises(CandidateKnowledgeConflictError, match="changed"):
        service.approve_document_selection(
            DocumentSelectionApproval(
                **request.model_dump(),
                document_version_id=platform.version.id,
                review_fingerprint="0" * 64,
                confirmation_phrase="SELECT REVIEWED DOCUMENT",
            )
        )
    audit = service.approve_document_selection(
        DocumentSelectionApproval(
            **request.model_dump(),
            document_version_id=platform.version.id,
            review_fingerprint=preview.review_fingerprint,
            confirmation_phrase="SELECT REVIEWED DOCUMENT",
        )
    )
    assert audit.document_version_id == platform.version.id
    assert audit.criteria["recommended_document_version_id"] == platform.version.id
    assert service.list_document_selection_audits(application.id) == [audit]
    selected = ApplicationRepository(session).get(application.id)
    assert selected is not None
    assert selected.selected_document_version_id == platform.version.id


def test_document_selection_exclusion_and_stable_tie_breaking(
    session: Session, tmp_path: Path
) -> None:
    profile_id = _profile(session)
    service = _service(session, tmp_path)
    first = service.import_document(
        profile_id,
        file_name="a.txt",
        data=b"Platform Engineer Kubernetes",
        kind=DocumentKind.RESUME,
        display_name="A",
        variant_label="Alpha",
        job_family_tags=["platform"],
        is_primary=False,
    )
    second = service.import_document(
        profile_id,
        file_name="b.txt",
        data=b"Platform Engineer Kubernetes",
        kind=DocumentKind.RESUME,
        display_name="B",
        variant_label="Beta",
        job_family_tags=["platform"],
        is_primary=False,
    )
    job = JobRepository(session).add(
        JobCreate(
            source="fixture",
            external_id="selection-2",
            employer="Stable Systems",
            title="Platform Engineer",
            description_hash=hashlib.sha256(b"stable-selection").hexdigest(),
        )
    )
    application = ApplicationRepository(session).add(
        ApplicationCreate(
            workflow_id="stable-selection-workflow",
            profile_id=profile_id,
            job_id=job.id,
        )
    )
    preview = service.preview_document_selection(
        DocumentSelectionRequest(application_id=application.id, prefer_primary=False)
    )
    assert [item.document_id for item in preview.recommendations] == [
        first.document.id,
        second.document.id,
    ]
    excluded = service.preview_document_selection(
        DocumentSelectionRequest(
            application_id=application.id,
            excluded_document_ids=[first.document.id],
            prefer_primary=False,
        )
    )
    assert excluded.recommended_document_version_id == second.version.id


def test_document_selection_does_not_count_partial_requirement_overlap(
    session: Session, tmp_path: Path
) -> None:
    profile_id = _profile(session)
    service = _service(session, tmp_path)
    service.import_document(
        profile_id,
        file_name="partial.txt",
        data=b"Kubernetes support",
        kind=DocumentKind.RESUME,
        display_name="Partial",
        variant_label="Partial",
        job_family_tags=[],
        is_primary=False,
    )
    job = JobRepository(session).add(
        JobCreate(
            source="fixture",
            external_id="selection-partial",
            employer="Conservative Systems",
            title="Platform Engineer",
            description_hash=hashlib.sha256(b"selection-partial").hexdigest(),
        )
    )
    session.add(
        JobRequirementRow(
            id="selection-kubernetes-operations",
            job_id=job.id,
            category="skill",
            text="Kubernetes operations",
            required=True,
            evidence_json={"source": "fixture"},
        )
    )
    session.commit()
    application = ApplicationRepository(session).add(
        ApplicationCreate(
            workflow_id="selection-partial-workflow",
            profile_id=profile_id,
            job_id=job.id,
        )
    )

    preview = service.preview_document_selection(
        DocumentSelectionRequest(application_id=application.id)
    )
    assert preview.recommendations[0].matched_requirement_ids == []
    assert any("0 of 1" in reason for reason in preview.recommendations[0].reasons)


def test_importing_primary_resume_replaces_prior_primary(session: Session, tmp_path: Path) -> None:
    profile_id = _profile(session)
    service = _service(session, tmp_path)
    first = service.import_document(
        profile_id,
        file_name="first.txt",
        data=b"First platform resume",
        kind=DocumentKind.RESUME,
        display_name="First",
        variant_label="First",
        job_family_tags=["platform"],
        is_primary=True,
    )
    second = service.import_document(
        profile_id,
        file_name="second.txt",
        data=b"Second platform resume",
        kind=DocumentKind.RESUME,
        display_name="Second",
        variant_label="Second",
        job_family_tags=["platform"],
        is_primary=True,
    )

    documents = {item.id: item for item in service.list_documents(profile_id)}
    assert documents[first.document.id].is_primary is False
    assert documents[second.document.id].is_primary is True


def test_pdf_docx_and_rtf_layout_extractors() -> None:
    pdf_buffer = BytesIO()
    canvas = Canvas(pdf_buffer)
    canvas.drawString(72, 720, "PDF resume with Python and BGP")
    canvas.save()
    _, pdf = extract_document("resume.pdf", pdf_buffer.getvalue())
    assert "PDF resume" in pdf.plain_text
    assert pdf.page_count == 1
    assert pdf.blocks[0].page == 1

    docx_buffer = BytesIO()
    document = DocxDocument()
    document.add_heading("DOCX Resume", level=1)
    document.add_paragraph("Python automation engineer")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Certification"
    table.cell(0, 1).text = "CCNP"
    document.save(docx_buffer)
    _, docx = extract_document("resume.docx", docx_buffer.getvalue())
    assert any(block.style == "Heading 1" for block in docx.blocks)
    assert any(block.kind == "TABLE_ROW" and "CCNP" in block.text for block in docx.blocks)

    _, rtf = extract_document(
        "resume.rtf", b"{\\rtf1\\ansi Candidate\\par Python automation\\par CCNP}"
    )
    assert "Python automation" in rtf.plain_text


def test_candidate_knowledge_multipart_api(session: Session, tmp_path: Path) -> None:
    profile_id = _profile(session)
    service = _service(session, tmp_path)
    app.dependency_overrides[get_knowledge_service] = lambda: service
    client = TestClient(app)
    try:
        response = client.post(
            f"/api/v1/knowledge/profiles/{profile_id}/documents",
            files={"file": ("resume.txt", RESUME, "text/plain")},
            data={
                "kind": "RESUME",
                "display_name": "API resume",
                "variant_label": "Infrastructure",
                "job_family_tags": "network,platform",
                "is_primary": "true",
            },
        )
        assert response.status_code == 201
        payload = response.json()
        assert payload["document"]["is_primary"] is True
        assert payload["version"]["media_type"] == "text/plain"
        assert payload["proposed_claims"]
        claim_id = payload["proposed_claims"][0]["id"]
        service.review_claim(
            claim_id,
            ClaimReview(approved=True, permitted_use=ClaimPermittedUse.APPLICATIONS),
        )
        created_answer = client.post(
            f"/api/v1/knowledge/profiles/{profile_id}/answers",
            json={
                "question": "Which verified skill should be reused?",
                "canonical_field": "skills.primary",
                "answer": "Use the verified skill from this profile.",
                "evidence_claim_ids": [claim_id],
                "confirmation_phrase": "SAVE REVIEWED ANSWER",
            },
        )
        assert created_answer.status_code == 201
        answer_payload = created_answer.json()
        updated_answer = client.put(
            f"/api/v1/knowledge/answers/{answer_payload['id']}",
            json={
                "question": answer_payload["question"],
                "canonical_field": answer_payload["canonical_field"],
                "answer": "Use only the verified skill from this profile.",
                "evidence_claim_ids": [claim_id],
                "expected_revision": 1,
                "confirmation_phrase": "SAVE REVIEWED ANSWER",
            },
        )
        assert updated_answer.status_code == 200
        assert updated_answer.json()["revision"] == 2
        history = client.get(f"/api/v1/knowledge/answers/{answer_payload['id']}/revisions")
        assert history.status_code == 200
        assert [item["revision"] for item in history.json()] == [2, 1]
        snapshot = client.get(f"/api/v1/knowledge/profiles/{profile_id}/snapshot").json()
        assert len(snapshot["documents"]) == 1
        assert snapshot["answers"][0]["revision"] == 2
    finally:
        app.dependency_overrides.clear()


def test_document_selection_api_requires_fingerprint_approval(
    session: Session, tmp_path: Path
) -> None:
    profile_id = _profile(session)
    service = _service(session, tmp_path)
    imported = service.import_document(
        profile_id,
        file_name="platform.txt",
        data=b"Senior Platform Engineer with Kubernetes and Python experience",
        kind=DocumentKind.RESUME,
        display_name="Platform resume",
        variant_label="Platform",
        job_family_tags=["platform", "cloud"],
        is_primary=True,
    )
    job = JobRepository(session).add(
        JobCreate(
            source="fixture",
            external_id="selection-api",
            employer="API Systems",
            title="Platform Engineer",
            description_hash=hashlib.sha256(b"selection-api").hexdigest(),
        )
    )
    application = ApplicationRepository(session).add(
        ApplicationCreate(
            workflow_id="selection-api-workflow",
            profile_id=profile_id,
            job_id=job.id,
        )
    )
    app.dependency_overrides[get_knowledge_service] = lambda: service
    client = TestClient(app)
    request = {
        "application_id": application.id,
        "kind": "RESUME",
        "preferred_tags": ["platform"],
        "excluded_document_ids": [],
        "prefer_primary": True,
    }
    try:
        preview_response = client.post(
            "/api/v1/knowledge/documents/selection/preview", json=request
        )
        assert preview_response.status_code == 200
        preview = preview_response.json()
        assert preview["recommended_document_version_id"] == imported.version.id

        stale = client.post(
            "/api/v1/knowledge/documents/selection/approve",
            json={
                **request,
                "document_version_id": imported.version.id,
                "review_fingerprint": "0" * 64,
                "confirmation_phrase": "SELECT REVIEWED DOCUMENT",
            },
        )
        assert stale.status_code == 409

        approved = client.post(
            "/api/v1/knowledge/documents/selection/approve",
            json={
                **request,
                "document_version_id": imported.version.id,
                "review_fingerprint": preview["review_fingerprint"],
                "confirmation_phrase": "SELECT REVIEWED DOCUMENT",
            },
        )
        assert approved.status_code == 200
        assert approved.json()["document_version_id"] == imported.version.id
        audits = client.get(f"/api/v1/knowledge/applications/{application.id}/document-selections")
        assert audits.status_code == 200
        assert [item["id"] for item in audits.json()] == [approved.json()["id"]]
    finally:
        app.dependency_overrides.clear()


def test_tailored_document_requires_locked_evidence_and_exact_review(
    session: Session, tmp_path: Path
) -> None:
    profile_id = _profile(session)
    service = _service(session, tmp_path)
    imported = service.import_document(
        profile_id,
        file_name="source-resume.txt",
        data=RESUME,
        kind=DocumentKind.RESUME,
        display_name="Source resume",
        variant_label="Source",
        job_family_tags=["network"],
        is_primary=True,
    )
    python_claim = next(
        claim for claim in imported.proposed_claims if claim.canonical_key == "skill.python"
    )
    service.review_claim(
        python_claim.id,
        ClaimReview(approved=True, permitted_use=ClaimPermittedUse.APPLICATIONS),
    )
    job = JobRepository(session).add(
        JobCreate(
            source="fixture",
            external_id="tailored-1",
            employer="Evidence Systems",
            title="Python Automation Engineer",
            description_hash=hashlib.sha256(b"python automation").hexdigest(),
        )
    )
    session.add(
        JobRequirementRow(
            id="requirement-python",
            job_id=job.id,
            category="skill",
            text="Production Python experience",
            required=True,
            evidence_json={"source": "fixture"},
        )
    )
    session.commit()
    application = ApplicationRepository(session).add(
        ApplicationCreate(
            workflow_id="tailored-workflow",
            profile_id=profile_id,
            job_id=job.id,
        )
    )
    request = TailoredDocumentRequest(
        application_id=application.id,
        kind=DocumentKind.RESUME,
        output_format=DocumentOutputFormat.DOCX,
        variant_label="Python evidence",
        template=DocumentTemplate.COMPACT,
    )
    preview = service.preview_tailored_document(request)
    assert preview.selected_claim_ids == [python_claim.id]
    assert preview.matched_requirement_ids == ["requirement-python"]
    assert preview.missing_required_requirements == []
    with pytest.raises(CandidateKnowledgeConflictError, match="changed"):
        service.generate_tailored_document(
            TailoredDocumentApproval(
                **request.model_dump(),
                review_fingerprint="0" * 64,
                confirmation_phrase="APPROVE TAILORED DOCUMENT",
            )
        )

    generated = service.generate_tailored_document(
        TailoredDocumentApproval(
            **request.model_dump(),
            review_fingerprint=preview.review_fingerprint,
            confirmation_phrase="APPROVE TAILORED DOCUMENT",
        )
    )
    assert generated.version.file_name.endswith(".docx")
    assert generated.document.variant_label == "Python evidence"
    assert "Python" in service.get_extraction(generated.version.id).plain_text
    assert service.get_document_content(generated.version.id).startswith(b"PK")
    stored = session.get(DocumentVersionRow, generated.version.id)
    assert stored is not None
    assert "Python" not in Path(stored.storage_path).read_text(encoding="ascii")
    audit = session.scalar(
        select(DocumentGenerationAuditRow).where(
            DocumentGenerationAuditRow.document_version_id == generated.version.id
        )
    )
    assert audit is not None
    assert audit.review_fingerprint == preview.review_fingerprint
    assert audit.template == "COMPACT"
    assert audit.ranking_method == "DETERMINISTIC_TOKEN_OVERLAP"
    rendered = DocxDocument(BytesIO(service.get_document_content(generated.version.id)))
    top_margin = rendered.sections[0].top_margin
    assert top_margin is not None
    assert top_margin.inches == pytest.approx(0.45, abs=0.01)
    assert service.list_generation_audits(application.id)[0] == generated.audit

    with pytest.raises(CandidateKnowledgeConflictError, match="hash changed"):
        service.capture_submitted_document(
            application.id,
            SubmittedDocumentCapture(
                document_version_id=generated.version.id,
                role=ApplicationDocumentRole.RESUME,
                expected_sha256="0" * 64,
                displayed_file_name=generated.version.file_name,
                upload_fingerprint="reference-ats:resume-upload-1",
                confirmation_phrase="RETAIN SUBMITTED DOCUMENT",
            ),
        )
    with pytest.raises(CandidateKnowledgeConflictError, match="filename"):
        service.capture_submitted_document(
            application.id,
            SubmittedDocumentCapture(
                document_version_id=generated.version.id,
                role=ApplicationDocumentRole.RESUME,
                expected_sha256=generated.version.sha256,
                displayed_file_name="wrong-file-name.docx",
                upload_fingerprint="reference-ats:resume-upload-1",
                confirmation_phrase="RETAIN SUBMITTED DOCUMENT",
            ),
        )
    capture = SubmittedDocumentCapture(
        document_version_id=generated.version.id,
        role=ApplicationDocumentRole.RESUME,
        expected_sha256=generated.version.sha256,
        displayed_file_name=generated.version.file_name,
        upload_fingerprint="reference-ats:resume-upload-1",
        confirmation_phrase="RETAIN SUBMITTED DOCUMENT",
    )
    submitted = service.capture_submitted_document(application.id, capture)
    assert submitted.sha256 == generated.version.sha256
    assert service.capture_submitted_document(application.id, capture) == submitted
    assert service.list_submitted_documents(application.id) == [submitted]

    cover_request = TailoredDocumentRequest(
        application_id=application.id,
        kind=DocumentKind.COVER_LETTER,
        output_format=DocumentOutputFormat.PDF,
        variant_label="Python cover letter",
    )
    cover_preview = service.preview_tailored_document(cover_request)
    cover = service.generate_tailored_document(
        TailoredDocumentApproval(
            **cover_request.model_dump(),
            review_fingerprint=cover_preview.review_fingerprint,
            confirmation_phrase="APPROVE TAILORED DOCUMENT",
        )
    )
    assert cover.version.file_name.endswith(".pdf")
    assert "Python Automation Engineer" in service.get_extraction(cover.version.id).plain_text


def test_tailored_document_governed_ranking_and_fallback_are_fingerprinted(
    session: Session, tmp_path: Path
) -> None:
    profile_id = _profile(session)
    base = _service(session, tmp_path)
    imported = base.import_document(
        profile_id,
        file_name="ranking.txt",
        data=b"Python\nKubernetes",
        kind=DocumentKind.RESUME,
        display_name="Ranking source",
        variant_label="Source",
        job_family_tags=["platform"],
        is_primary=True,
    )
    for claim in imported.proposed_claims:
        base.review_claim(
            claim.id,
            ClaimReview(approved=True, permitted_use=ClaimPermittedUse.APPLICATIONS),
        )
    job = JobRepository(session).add(
        JobCreate(
            source="fixture",
            external_id="ranking-1",
            employer="Evidence Systems",
            title="Platform Engineer",
            description_hash=hashlib.sha256(b"platform").hexdigest(),
        )
    )
    session.add(
        JobRequirementRow(
            id="requirement-kubernetes",
            job_id=job.id,
            category="skill",
            text="Kubernetes operations",
            required=True,
            evidence_json={"source": "fixture"},
        )
    )
    session.commit()
    application = ApplicationRepository(session).add(
        ApplicationCreate(
            workflow_id="ranking-workflow",
            profile_id=profile_id,
            job_id=job.id,
        )
    )

    class ReversingReranker:
        def rerank(self, request: AIRerankRequest) -> list[AIRerankResult]:
            assert request.classification is DataClassification.EMPLOYMENT_SENSITIVE
            assert request.external_consent is False
            return [AIRerankResult(index=1, score=0.9), AIRerankResult(index=0, score=0.8)]

    ranked_service = CandidateKnowledgeService(
        CandidateKnowledgeRepository(session),
        CandidateRepository(session),
        JobRepository(session),
        ApplicationRepository(session),
        SensitiveDataCipher(StaticKeyProvider(b"k" * 32)),
        document_data_dir=tmp_path / "ranked-documents",
        document_max_bytes=10_485_760,
        ai_gateway=ReversingReranker(),
    )
    request = TailoredDocumentRequest(
        application_id=application.id,
        kind=DocumentKind.RESUME,
        template=DocumentTemplate.COMPACT,
        ranking_mode=TailoringRankingMode.GOVERNED_AI,
        max_claims=2,
    )
    preview = ranked_service.preview_tailored_document(request)
    assert preview.template is DocumentTemplate.COMPACT
    assert preview.ranking_method == "GOVERNED_AI_RERANK"
    assert preview.ranking_notice is None
    deterministic_ids = [
        claim.id
        for claim in sorted(
            imported.proposed_claims,
            key=lambda claim: (claim.canonical_key, claim.id),
        )
    ]
    assert preview.selected_claim_ids == list(reversed(deterministic_ids))

    fallback = base.preview_tailored_document(request)
    assert fallback.ranking_method == "DETERMINISTIC_FALLBACK"
    assert fallback.ranking_notice is not None
    assert fallback.review_fingerprint != preview.review_fingerprint
