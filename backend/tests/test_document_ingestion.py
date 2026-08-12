from __future__ import annotations

import subprocess
import zipfile
from io import BytesIO
from pathlib import Path
from typing import cast

import pytest
from docx import Document as DocxDocument
from PIL import Image, ImageDraw
from pypdf import PdfWriter
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen.canvas import Canvas

from job_apply_pro.documents import extractors
from job_apply_pro.documents.extractors import (
    DocumentExtractionError,
    DocumentIngestionOptions,
    UnsupportedDocumentError,
    extract_document,
)


def _docx_bytes(text: str) -> bytes:
    output = BytesIO()
    document = DocxDocument()
    document.add_paragraph(text)
    document.save(output)
    return output.getvalue()


def _image_only_pdf(text: str) -> bytes:
    image = Image.new("RGB", (900, 160), "white")
    ImageDraw.Draw(image).text((20, 50), text, fill="black")
    output = BytesIO()
    canvas = Canvas(output, pagesize=(900, 160))
    canvas.drawImage(ImageReader(image), 0, 0, width=900, height=160)
    canvas.save()
    image.close()
    return output.getvalue()


def test_legacy_doc_requires_explicit_trusted_converter() -> None:
    with pytest.raises(UnsupportedDocumentError, match="explicitly configured"):
        extract_document("resume.doc", b"legacy-binary-data")


def test_legacy_doc_conversion_uses_fixed_shell_free_command(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    executable = tmp_path / "soffice.exe"
    executable.write_bytes(b"fixture")
    captured: list[str] = []

    def fake_run(
        command: list[str], *, cwd: Path, timeout: int
    ) -> subprocess.CompletedProcess[str]:
        captured.extend(command)
        assert timeout == 17
        (cwd / "source.docx").write_bytes(_docx_bytes("Converted legacy resume"))
        return subprocess.CompletedProcess(command, 0, "converted", "")

    monkeypatch.setattr(extractors, "_run_process", fake_run)
    media_type, extraction = extract_document(
        "resume.doc",
        b"legacy-binary-data",
        DocumentIngestionOptions(
            legacy_doc_converter_path=executable.resolve(),
            legacy_doc_conversion_timeout_seconds=17,
        ),
    )

    assert media_type == "application/msword"
    assert extraction.plain_text == "Converted legacy resume"
    assert extraction.parser == "libreoffice-doc-to-docx/1+python-docx-layout/2"
    assert extraction.warnings
    assert captured[0] == str(executable.resolve())
    assert "--headless" in captured
    assert "--convert-to" in captured


def test_legacy_doc_rejects_unapproved_executable_name(tmp_path: Path) -> None:
    executable = tmp_path / "cmd.exe"
    executable.write_bytes(b"fixture")
    with pytest.raises(DocumentExtractionError, match="approved executable"):
        extract_document(
            "resume.doc",
            b"legacy-binary-data",
            DocumentIngestionOptions(legacy_doc_converter_path=executable.resolve()),
        )


def test_legacy_doc_rejects_oversized_converted_output(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    executable = tmp_path / "soffice.exe"
    executable.write_bytes(b"fixture")

    def fake_run(
        command: list[str], *, cwd: Path, timeout: int
    ) -> subprocess.CompletedProcess[str]:
        (cwd / "source.docx").write_bytes(b"x" * 100)
        return subprocess.CompletedProcess(command, 0, "converted", "")

    monkeypatch.setattr(extractors, "_run_process", fake_run)
    with pytest.raises(DocumentExtractionError, match="exceeded safety limits"):
        extract_document(
            "resume.doc",
            b"legacy-binary-data",
            DocumentIngestionOptions(
                legacy_doc_converter_path=executable.resolve(), max_converted_bytes=10
            ),
        )


def test_image_only_pdf_has_actionable_ocr_disabled_error() -> None:
    with pytest.raises(DocumentExtractionError, match="configure and enable OCR"):
        extract_document("scanned.pdf", _image_only_pdf("Scanned candidate resume"))


def test_docx_rejects_oversized_expanded_entry(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(extractors, "MAX_DOCX_SINGLE_ENTRY_BYTES", 10)
    output = BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("word/document.xml", "x" * 100)
    with pytest.raises(DocumentExtractionError, match="entry exceeds"):
        extract_document("oversized.docx", output.getvalue())


def test_malformed_docx_and_pdf_are_normalized() -> None:
    with pytest.raises(DocumentExtractionError, match="DOCX archive validation failed"):
        extract_document("malformed.docx", b"not-a-zip")
    with pytest.raises(DocumentExtractionError, match="PDF parsing failed"):
        extract_document("malformed.pdf", b"not-a-pdf")


def test_password_protected_pdf_is_rejected() -> None:
    output = BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.encrypt("secret")
    writer.write(output)
    with pytest.raises(DocumentExtractionError, match="Password-protected"):
        extract_document("protected.pdf", output.getvalue())


def test_partially_textless_pdf_is_imported_with_warning() -> None:
    output = BytesIO()
    canvas = Canvas(output)
    canvas.drawString(72, 720, "Candidate resume text")
    canvas.showPage()
    canvas.showPage()
    canvas.save()
    _, extraction = extract_document("partial.pdf", output.getvalue())
    assert extraction.plain_text == "Candidate resume text"
    assert extraction.warnings == ["1 PDF page(s) had no meaningful text; OCR was not enabled"]


def test_pdf_layout_orders_two_columns_for_resume_reading() -> None:
    output = BytesIO()
    canvas = Canvas(output)
    canvas.drawString(72, 740, "Candidate Name")
    canvas.drawString(72, 700, "Experience")
    canvas.drawString(330, 700, "Skills")
    canvas.drawString(72, 680, "Platform Engineer")
    canvas.drawString(330, 680, "Python")
    canvas.save()

    _, extraction = extract_document("two-column.pdf", output.getvalue())

    assert extraction.parser == "pypdf-layout/2"
    assert [block.text for block in extraction.blocks] == [
        "Candidate Name",
        "Experience",
        "Platform Engineer",
        "Skills",
        "Python",
    ]
    assert [block.column for block in extraction.blocks] == [None, 0, 0, 1, 1]
    assert extraction.warnings == [
        "Layout-aware column ordering was applied to PDF page(s): 1; "
        "review complex graphics or spanning content"
    ]


def test_pdf_layout_does_not_promote_one_off_indented_content_to_a_column() -> None:
    output = BytesIO()
    canvas = Canvas(output)
    canvas.drawString(72, 740, "Candidate Name")
    canvas.drawString(72, 700, "Experience")
    canvas.drawString(330, 700, "2021 to present")
    canvas.drawString(72, 680, "Platform Engineer")
    canvas.save()

    _, extraction = extract_document("indented.pdf", output.getvalue())

    assert [block.text for block in extraction.blocks] == [
        "Candidate Name",
        "Experience 2021 to present",
        "Platform Engineer",
    ]
    assert all(block.column is None for block in extraction.blocks)
    assert extraction.warnings == []


def test_docx_preserves_paragraph_and_table_document_order() -> None:
    output = BytesIO()
    document = DocxDocument()
    document.add_paragraph("Summary")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Skill"
    table.cell(0, 1).text = "Python"
    document.add_paragraph("Experience")
    document.save(output)

    _, extraction = extract_document("mixed-layout.docx", output.getvalue())

    assert extraction.parser == "python-docx-layout/2"
    assert [block.text for block in extraction.blocks] == [
        "Summary",
        "Skill | Python",
        "Experience",
    ]
    assert extraction.blocks[1].table == 0
    assert extraction.blocks[1].row == 0


def test_ocr_rejects_unapproved_executable_name(tmp_path: Path) -> None:
    executable = tmp_path / "cmd.exe"
    executable.write_bytes(b"fixture")
    with pytest.raises(DocumentExtractionError, match="approved executable"):
        extract_document(
            "scan.pdf",
            _image_only_pdf("Scanned candidate resume"),
            DocumentIngestionOptions(ocr_enabled=True, ocr_tesseract_path=executable.resolve()),
        )


def test_image_only_pdf_uses_bounded_tesseract_fallback(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    executable = tmp_path / "tesseract.exe"
    executable.write_bytes(b"fixture")
    captured: list[str] = []

    def fake_run(
        command: list[str], *, cwd: Path, timeout: int
    ) -> subprocess.CompletedProcess[str]:
        captured.extend(command)
        assert cwd.name.startswith("jap-ocr-")
        assert timeout == 11
        assert Path(command[1]).is_file()
        Path(f"{command[2]}.txt").write_text("OCR recovered candidate resume", encoding="utf-8")
        return subprocess.CompletedProcess(command, 0, "", "")

    monkeypatch.setattr(extractors, "_run_process", fake_run)
    _, extraction = extract_document(
        "scanned.pdf",
        _image_only_pdf("Scanned candidate resume"),
        DocumentIngestionOptions(
            ocr_enabled=True,
            ocr_tesseract_path=executable.resolve(),
            ocr_language="eng",
            ocr_dpi=150,
            ocr_page_timeout_seconds=11,
        ),
    )

    assert extraction.plain_text == "OCR recovered candidate resume"
    assert extraction.parser == "pypdf-layout/2+tesseract/1"
    assert extraction.blocks[0].kind == "OCR_PAGE_TEXT"
    assert extraction.blocks[0].style == "tesseract:eng:150dpi"
    assert captured[0] == str(executable.resolve())
    assert captured[-4:] == ["-l", "eng", "--psm", "6"]


def test_process_timeout_is_normalized(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    def timeout(*args: object, **kwargs: object) -> subprocess.CompletedProcess[str]:
        raise subprocess.TimeoutExpired(cmd="fixture", timeout=1)

    monkeypatch.setattr(subprocess, "run", timeout)
    with pytest.raises(DocumentExtractionError, match="timed out"):
        extractors._run_process(["fixture"], cwd=tmp_path, timeout=1)


def test_process_environment_excludes_backend_secrets(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    captured: dict[str, str] = {}

    def succeed(*args: object, **kwargs: object) -> subprocess.CompletedProcess[str]:
        captured.update(cast(dict[str, str], kwargs["env"]))
        return subprocess.CompletedProcess(["fixture"], 0, "", "")

    monkeypatch.setenv("JAP_MASTER_KEY", "must-not-cross-process-boundary")
    monkeypatch.setenv("JAP_API_TOKEN", "must-not-cross-process-boundary")
    monkeypatch.setattr(subprocess, "run", succeed)
    extractors._run_process(["fixture"], cwd=tmp_path, timeout=1)

    assert "JAP_MASTER_KEY" not in captured
    assert "JAP_API_TOKEN" not in captured
    assert captured["TEMP"] == str(tmp_path)
    assert captured["USERPROFILE"] == str(tmp_path)
